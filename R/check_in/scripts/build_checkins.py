#!/usr/bin/env python3
"""Build check-in quiz Word documents (blank / solution / rubric) from JSON.

Usage:
  python R/check_in/scripts/build_checkins.py --week 1
  python R/check_in/scripts/build_checkins.py --week 1 --type self
  python R/check_in/scripts/build_checkins.py --week 1 --type collab
  python R/check_in/scripts/build_checkins.py --all
  python R/check_in/scripts/build_checkins.py --week 1 --versions blank solution

Point values are read ONLY from the JSON ("points" on each question). A null
value renders as "(__ PTS)" and produces a warning so you remember to set it.
"""
import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, Inches, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("python-docx is required:  pip install python-docx")

ROOT = Path(__file__).resolve().parents[1]            # R/check_in
SOURCE_DIRS = {"self": ROOT / "check_in_self", "collab": ROOT / "check_in_collab"}
GENERATED = ROOT / "generated"
FORMAT_FILE = ROOT / "templates" / "format.json"
VERSIONS = ("blank", "solution", "rubric")
VALID_TYPES = tuple(SOURCE_DIRS)
VALID_QTYPES = {"short_answer", "multiple_choice", "true_false", "code"}


# --------------------------------------------------------------------------- #
# Validation
# --------------------------------------------------------------------------- #
class ValidationError(Exception):
    pass


def _is_number(x):
    return isinstance(x, (int, float)) and not isinstance(x, bool)


def validate(quiz, path):
    errs, warns = [], []
    where = path.name

    for key in ("week", "check_in_type", "title", "questions"):
        if key not in quiz:
            errs.append(f"{where}: missing top-level key '{key}'")
    if errs:
        raise ValidationError("\n".join(errs))

    if quiz["check_in_type"] not in VALID_TYPES:
        errs.append(f"{where}: check_in_type must be one of {VALID_TYPES}, got '{quiz['check_in_type']}'")
    if not _is_number(quiz["week"]) or quiz["week"] < 0:
        errs.append(f"{where}: 'week' must be a non-negative number")
    for flag in ("blank_published", "solution_published"):
        if flag in quiz and not isinstance(quiz[flag], bool):
            errs.append(f"{where}: '{flag}' must be true or false")

    qs = quiz["questions"]
    if not isinstance(qs, list) or len(qs) == 0:
        errs.append(f"{where}: 'questions' must be a non-empty list")
        raise ValidationError("\n".join(errs))

    seen = set()
    for i, q in enumerate(qs, 1):
        tag = f"{where} question #{i}"
        qid = q.get("id")
        if not qid or not isinstance(qid, str):
            errs.append(f"{tag}: missing 'id'")
        elif qid in seen:
            errs.append(f"{tag}: duplicate id '{qid}'")
        else:
            seen.add(qid)
            tag = f"{where} [{qid}]"
        if not str(q.get("question", "")).strip():
            errs.append(f"{tag}: missing 'question' text")
        if not str(q.get("answer", "")).strip():
            errs.append(f"{tag}: missing 'answer'")
        if not str(q.get("rubric", "")).strip():
            errs.append(f"{tag}: missing 'rubric'")
        pts = q.get("points", "MISSING")
        if pts == "MISSING":
            errs.append(f"{tag}: missing 'points' (use null if undecided)")
        elif pts is None:
            warns.append(f"{tag}: points not set (null) — will render as '__'")
        elif not _is_number(pts) or pts < 0:
            errs.append(f"{tag}: 'points' must be null or a non-negative number, got {pts!r}")
        if "type" in q and q["type"] not in VALID_QTYPES:
            errs.append(f"{tag}: unknown question type '{q['type']}' (valid: {sorted(VALID_QTYPES)})")
        if "extra_credit" in q and not isinstance(q["extra_credit"], bool):
            errs.append(f"{tag}: 'extra_credit' must be true/false")
        if "[INSTRUCTOR" in str(q.get("answer", "")):
            warns.append(f"{tag}: answer still contains an [INSTRUCTOR ...] placeholder")

    if errs:
        raise ValidationError("\n".join(errs))
    return warns


# --------------------------------------------------------------------------- #
# Word helpers
# --------------------------------------------------------------------------- #
def fmt_points(p, fmt):
    if p is None:
        return fmt["points_placeholder"]
    return str(int(p)) if float(p).is_integer() else f"{p:g}"


def totals(quiz, fmt):
    reg = [q["points"] for q in quiz["questions"] if not q.get("extra_credit")]
    ec = [q["points"] for q in quiz["questions"] if q.get("extra_credit")]
    t = fmt["points_placeholder"] if any(p is None for p in reg) else fmt_points(sum(reg), fmt)
    e = None
    if ec:
        e = fmt["points_placeholder"] if any(p is None for p in ec) else fmt_points(sum(ec), fmt)
    return t, e


def add_runs(par, text, bold=False, italic=False, font=None, size=None, color=None):
    """Add text; supports **bold** spans inside the string."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_b = part.startswith("**") and part.endswith("**")
        run = par.add_run(part[2:-2] if is_b else part)
        run.bold = bold or is_b
        run.italic = italic
        if font:
            run.font.name = font
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(attr), font)
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return par


def new_document(fmt, header_line=None):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = fmt["font_name"]
    st.font.size = Pt(fmt["font_size_pt"])
    st.element.rPr.rFonts.set(qn("w:eastAsia"), fmt["font_name"])
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.space_before = Pt(0)
    sec = doc.sections[0]
    m = fmt["margins_inches"]
    sec.top_margin, sec.bottom_margin = Inches(m["top"]), Inches(m["bottom"])
    sec.left_margin, sec.right_margin = Inches(m["left"]), Inches(m["right"])
    # page header: Name / username / date on every page (as in the spec)
    hp = sec.header.paragraphs[0]
    hp.text = ""
    add_runs(hp, header_line or fmt["header_line"], font=fmt["font_name"], size=fmt["font_size_pt"])
    return doc


def para(doc, text="", bold=False, italic=False, align=None, fmt=None, size=None, space_after=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        add_runs(p, text, bold=bold, italic=italic, font=fmt["font_name"], size=size or fmt["font_size_pt"])
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def code_block(doc, text, fmt):
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(fmt["font_size_pt"] - 1)
        r._element.get_or_add_rPr()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def question_text(doc, q, fmt):
    """Question text; lines indented by 4 spaces are rendered as code."""
    blocks = re.split(r"\n\n", q["question"].strip())
    first = True
    for b in blocks:
        if b.startswith("    "):
            code_block(doc, "\n".join(l[4:] if l.startswith("    ") else l for l in b.splitlines()), fmt)
        else:
            p = doc.add_paragraph()
            if first:
                label = f"{fmt['extra_credit_prefix'] if q.get('extra_credit') else ''}Q{q['_n']}. ({fmt_points(q['points'], fmt)} {fmt['points_label']}) "
                add_runs(p, label, bold=True, font=fmt["font_name"], size=fmt["font_size_pt"])
            add_runs(p, b, font=fmt["font_name"], size=fmt["font_size_pt"])
        first = False


def answer_space(doc, inches):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(int(inches * 72))


def intro(doc, quiz, fmt, version):
    title = fmt["title_pattern"].format(title=quiz["title"], version=quiz.get("version", "A"))
    para(doc, title, bold=True, align="center", fmt=fmt, size=fmt["title_size_pt"], space_after=4)
    if version == "solution":
        para(doc, fmt["solution_banner"], bold=True, align="center", fmt=fmt, space_after=6)
    elif version == "rubric":
        para(doc, fmt["rubric_banner"], bold=True, align="center", fmt=fmt, space_after=6)
    text = quiz.get("instructions", "").replace("{time_limit}", f"**{quiz.get('time_limit', '')}**")
    t, e = totals(quiz, fmt)
    sentences = [text, fmt["total_sentence"].format(total=f"**{t}**")]
    if e is not None:
        sentences.append(fmt["extra_credit_sentence"].format(ec_total=f"**{e}**"))
    para(doc, " ".join(s for s in sentences if s), fmt=fmt, space_after=10)


def number_questions(quiz):
    for n, q in enumerate(quiz["questions"], 1):
        q["_n"] = n


# --------------------------------------------------------------------------- #
# The three versions
# --------------------------------------------------------------------------- #
def build_blank(quiz, fmt, doc=None):
    doc = doc or new_document(fmt, quiz.get("header_line"))
    intro(doc, quiz, fmt, "blank")
    for q in quiz["questions"]:
        question_text(doc, q, fmt)
        answer_space(doc, q.get("answer_space_inches", fmt["default_answer_space_inches"]))
    return doc


def build_solution(quiz, fmt, doc=None):
    doc = doc or new_document(fmt, quiz.get("header_line"))
    intro(doc, quiz, fmt, "solution")
    for q in quiz["questions"]:
        question_text(doc, q, fmt)
        ink = fmt.get("solution_ink_hex")  # red ink for everything the instructor adds
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        add_runs(p, "Model answer: ", bold=True, font=fmt["font_name"], size=fmt["font_size_pt"], color=ink)
        add_runs(p, q["answer"], font=fmt["font_name"], size=fmt["font_size_pt"], color=ink)
        if q.get("instructor_notes"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, "Instructor notes: ", bold=True, italic=True, font=fmt["font_name"], size=fmt["font_size_pt"], color=ink)
            add_runs(p, q["instructor_notes"], italic=True, font=fmt["font_name"], size=fmt["font_size_pt"], color=ink)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return doc


def build_rubric(quiz, fmt, doc=None):
    doc = doc or new_document(fmt, quiz.get("header_line"))
    intro(doc, quiz, fmt, "rubric")
    t, e = totals(quiz, fmt)
    para(doc, f"**Total:** {t} {fmt['points_label']}" + (f" · **Extra credit:** {e} {fmt['points_label']}" if e else ""), fmt=fmt, space_after=8)
    for q in quiz["questions"]:
        question_text(doc, q, fmt)
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        rows = [("Points", f"{fmt_points(q['points'], fmt)} {fmt['points_label']}" + ("  (extra credit)" if q.get("extra_credit") else "")),
                ("Model answer", q["answer"]),
                ("Grading criteria", q["rubric"])]
        for (label, val), row in zip(rows, table.rows):
            row.cells[0].width = Inches(1.3)
            row.cells[1].width = Inches(5.2)
            row.cells[0].paragraphs[0].text = ""
            add_runs(row.cells[0].paragraphs[0], label, bold=True, font=fmt["font_name"], size=fmt["font_size_pt"])
            row.cells[1].paragraphs[0].text = ""
            add_runs(row.cells[1].paragraphs[0], val, font=fmt["font_name"], size=fmt["font_size_pt"])
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return doc


BUILDERS = {"blank": build_blank, "solution": build_solution, "rubric": build_rubric}


# --------------------------------------------------------------------------- #
# Driver
# --------------------------------------------------------------------------- #
def load_quiz(path):
    try:
        quiz = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        raise ValidationError(f"{path}: invalid JSON — {e}")
    for w in validate(quiz, path):
        print(f"  warning: {w}")
    return quiz


def build_one(path, versions, fmt):
    quiz = load_quiz(path)
    number_questions(quiz)
    wk, ct = int(quiz["week"]), quiz["check_in_type"]
    out_dir = GENERATED / f"week_{wk:02d}" / ct
    out_dir.mkdir(parents=True, exist_ok=True)
    # week_02.json -> week_02_collab_blank.docx ; week_02_B.json -> week_02_collab_B_blank.docx
    m = re.fullmatch(r"week_\d{2}_([A-Z])", path.stem)
    suffix = f"_{m.group(1)}" if m else ""
    outputs = []
    for v in versions:
        out = out_dir / f"week_{wk:02d}_{ct}{suffix}_{v}.docx"
        BUILDERS[v](quiz, fmt).save(out)
        outputs.append(out)
        print(f"  wrote {out.relative_to(ROOT.parent.parent)}")
    return outputs


def build_combined(paths, versions, fmt):
    """Several versions (e.g. B and C) in ONE .docx per output type, page break between them."""
    quizzes = [load_quiz(p) for p in paths]
    for q in quizzes:
        number_questions(q)
    letters = "".join(q.get("version", "A") for q in quizzes)
    wk, ct = int(quizzes[0]["week"]), quizzes[0]["check_in_type"]
    out_dir = GENERATED / f"week_{wk:02d}" / ct
    out_dir.mkdir(parents=True, exist_ok=True)
    for v in versions:
        doc = None
        for i, q in enumerate(quizzes):
            if doc is not None:
                doc.add_page_break()
            doc = BUILDERS[v](q, fmt, doc)
        out = out_dir / f"week_{wk:02d}_{ct}_{letters}_{v}.docx"
        doc.save(out)
        print(f"  wrote {out.relative_to(ROOT.parent.parent)}")


def find_sources(week=None, ctype=None):
    types = [ctype] if ctype else list(VALID_TYPES)
    files = []
    for t in types:
        pattern = f"week_{week:02d}*.json" if week is not None else "week_*.json"
        files += sorted(SOURCE_DIRS[t].glob(pattern))  # week_NN.json (A) and week_NN_B.json, week_NN_C.json
    return files


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--week", type=int, help="week number, e.g. 1")
    ap.add_argument("--type", choices=VALID_TYPES, help="self or collab (default: both)")
    ap.add_argument("--all", action="store_true", help="build every week found")
    ap.add_argument("--versions", nargs="+", choices=VERSIONS, default=list(VERSIONS),
                    help="which versions to build (default: all three)")
    ap.add_argument("--combine", nargs="+", metavar="LETTER",
                    help="build these versions into ONE document per output type, e.g. --combine B C")
    a = ap.parse_args(argv)
    if a.week is None and not a.all:
        ap.error("give --week N or --all")
    if a.type and a.type not in VALID_TYPES:
        ap.error(f"invalid check-in type '{a.type}'")

    fmt = json.loads(FORMAT_FILE.read_text(encoding="utf-8"))
    sources = find_sources(None if a.all else a.week, a.type)
    if not sources:
        sys.exit(f"No JSON found for week={a.week} type={a.type or 'both'} under {ROOT}")

    if a.combine:
        if a.week is None or not a.type:
            ap.error("--combine needs --week N and --type")
        paths = []
        for L in a.combine:
            p = SOURCE_DIRS[a.type] / (f"week_{a.week:02d}.json" if L == "A" else f"week_{a.week:02d}_{L}.json")
            if not p.exists():
                sys.exit(f"missing {p}")
            paths.append(p)
        print(f"combined {'+'.join(a.combine)}:")
        try:
            build_combined(paths, a.versions, fmt)
        except ValidationError as e:
            sys.exit(f"  ERROR — not built:\n    " + str(e).replace("\n", "\n    "))
        return

    failed = 0
    for src in sources:
        print(f"{src.relative_to(ROOT)}:")
        try:
            build_one(src, a.versions, fmt)
        except ValidationError as e:
            failed += 1
            print(f"  ERROR — not built:\n    " + str(e).replace("\n", "\n    "))
    if failed:
        sys.exit(f"{failed} check-in(s) failed validation")
    print("Done. Generated files live in R/check_in/generated/ (git-ignored, NOT on the website).")


if __name__ == "__main__":
    main()
